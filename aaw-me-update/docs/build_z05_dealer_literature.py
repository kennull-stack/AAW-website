from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "aaw-me-update" / "docs"
ASSETS = ROOT / "aaw-theme" / "assets"


BLACK = RGBColor(12, 12, 12)
BODY = RGBColor(42, 38, 32)
MUTED = RGBColor(100, 92, 82)
GOLD = RGBColor(177, 126, 43)
WHITE = RGBColor(255, 255, 255)
CREAM = RGBColor(246, 242, 234)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color="D8D0C3", size="5"):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:tcBorders"))
    if node is None:
        node = OxmlElement("w:tcBorders")
        tc_pr.append(node)
    for edge in ("top", "left", "bottom", "right"):
        el = node.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            node.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def margins(cell, top=120, start=150, bottom=120, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def table_width(table, width_dxa=10250):
    pr = table._tbl.tblPr
    w = pr.find(qn("w:tblW"))
    if w is None:
        w = OxmlElement("w:tblW")
        pr.append(w)
    w.set(qn("w:w"), str(width_dxa))
    w.set(qn("w:type"), "dxa")


def set_widths(table, widths):
    for row in table.rows:
        for i, width in enumerate(widths):
            cell = row.cells[i]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tcw = tc_pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tc_pr.append(tcw)
            tcw.set(qn("w:w"), str(int(width.inches * 1440)))
            tcw.set(qn("w:type"), "dxa")


def run(p, text, bold=False, size=9.4, color=BODY, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def setup(doc, zh=False):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.62)
    sec.right_margin = Inches(0.62)
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if zh else "Arial")
    doc.styles["Normal"].font.size = Pt(9.4)
    doc.styles["Normal"].font.color.rgb = BODY


def kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, text.upper(), bold=True, size=8, color=GOLD)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run(p, text, bold=True, size=22 if level == 1 else 14.5, color=BLACK)


def body(doc, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    run(p, text, size=9.35, color=BODY)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.26)
    p.paragraph_format.space_after = Pt(2)
    run(p, text, size=9.1, color=BODY)


def image(doc, filename, width=6.5):
    path = ASSETS / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(str(path), width=Inches(width))


def kv_table(doc, rows, widths=(1.75, 5.35), dark=False):
    t = doc.add_table(rows=len(rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    table_width(t)
    set_widths(t, [Inches(widths[0]), Inches(widths[1])])
    for i, (label, value) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        for c in (c0, c1):
            shade(c, "111111" if dark else "FFFFFF")
            borders(c, "3B3B3B" if dark else "D8D0C3")
            margins(c, 120, 150, 120, 150)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run(c0.paragraphs[0], label, bold=True, size=8.4, color=GOLD)
        run(c1.paragraphs[0], value, size=9, color=WHITE if dark else BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def matrix_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    table_width(t)
    if widths:
        set_widths(t, [Inches(w) for w in widths])
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, "111111")
        borders(cell, "111111")
        margins(cell, 120, 140, 120, 140)
        run(cell.paragraphs[0], header, bold=True, size=8.2, color=WHITE)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            shade(cell, "FFFFFF")
            borders(cell)
            margins(cell, 110, 135, 110, 135)
            run(cell.paragraphs[0], value, size=8.55, color=BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def cover(doc, data, zh=False):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    table_width(t)
    set_widths(t, [Inches(3.2), Inches(4.0)])
    for c in t.rows[0].cells:
        shade(c, "080808")
        borders(c, "080808", "0")
        margins(c, 260, 240, 260, 240)
    left, right = t.rows[0].cells
    run(left.paragraphs[0], data["cover_kicker"].upper(), bold=True, size=8, color=GOLD)
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    run(p, "Z05 AuAg", bold=True, size=29, color=WHITE)
    p = left.add_paragraph()
    run(p, data["launch"], bold=True, size=17, color=GOLD)
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run(p, data["intro"], size=9.3, color=CREAM)
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    run(p, data["slogan"], bold=True, size=9.1, color=CREAM)
    img = ASSETS / "z05-auag-luxury-portrait-hero.png"
    if img.exists():
        pr = right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr.add_run().add_picture(str(img), width=Inches(3.75))


EN = {
    "cover_kicker": "AAW Z Series Universal - Dealer Product Literature",
    "launch": "AuAg Launch Edition",
    "intro": "A compact universal hybrid for listeners who rotate music, sources, and moods. Z05 AuAg combines 10mm dynamic-driver punch, armature clarity, TwinCore physical voicing, hand-finished AuAg shell work, and a practical cable kit for daily use.",
    "slogan": "Those who listen will find the way.",
    "position": [
        "Z05 AuAg should be positioned as a small but serious universal-fit hybrid IEM: easy to carry, easy to drive, visually distinctive, and flexible enough for changing playlists.",
        "Compared with a one-signature demo monitor, Z05 AuAg is built around rotation: albums, playlists, phone dongles, DAPs, compact DAC/amps, desktop chains, and listening moods that change by the hour.",
        "The launch edition identity is AuAg: gold and silver accents over hand-finished shell work, with forged-carbon faceplates doing the visual heavy lifting."
    ],
    "selling": [
        ("Daily universal hybrid", "A compact universal shell with over-ear cable routing and LiquidMorph tips for repeatable seal and comfort."),
        ("TwinCore physical voicing", "A two-position switch on the shell changes overall voicing without tools, apps, or recessed micro-slots."),
        ("Dynamic-driver body", "A 10mm PET diaphragm dynamic woofer gives kick drums and bass lines real pressure and texture."),
        ("Armature clarity", "Mid and super-high BA drivers add vocal edge, cymbal air, and room cues without thinning the body."),
        ("Signal-chain ready kit", "2-pin detachable cable with 3.5mm stereo, 4.4mm balanced, and Type-C terminations in the box."),
    ],
    "twin": [
        "Blue Mode", "Weight / Flow", "Fuller low end, calmer upper harmonics, and an easier center image for long albums.",
        "Red Mode", "Focus / Spark", "Tighter bass outlines, more forward clarity, and extra cymbal air when the mix needs to open up."
    ],
    "specs": [
        ("Product", "AAW Z05 AuAg"),
        ("Format", "Universal-fit hybrid in-ear monitor"),
        ("Low-frequency driver", "10mm PET diaphragm dynamic woofer"),
        ("Balanced armature system", "Mid and super-high BA drivers for vocal edge and air"),
        ("Voicing control", "TwinCore physical voicing switch"),
        ("Pressure management", "Bass-Flow vent pressure system"),
        ("Shell body", "3D-printed universal monitor shell"),
        ("Faceplate", "Forged carbon with Au/Ag foil identity"),
        ("Finish", "Hand-laced gold and silver foil finish; pair-by-pair hand work"),
        ("Cable connection", "2-pin detachable monitor-side connection"),
        ("Source terminations", "3.5mm stereo, 4.4mm balanced, Type-C"),
        ("Sensitivity", "110dB @ 1kHz"),
        ("Impedance", "12 Ohm @ 1kHz"),
        ("THD", "<0.5% @ 1kHz"),
    ],
    "box": [
        ("Z05 AuAg monitors", "Left/right universal-fit hybrid IEMs with AuAg launch finish."),
        ("2-pin detachable cable", "Braided cable with familiar monitor-side 2-pin connection for long-term flexibility."),
        ("3.5mm stereo termination", "For standard analog outputs and broad compatibility."),
        ("4.4mm balanced termination", "For DAPs and compact DAC/amps with balanced output."),
        ("Type-C termination", "For direct use with compatible phones, tablets, laptops, and compact digital setups."),
        ("Carrying case", "Storage and travel protection."),
        ("LiquidMorph eartips", "Clear flange and teal bore; supports seal stability, bass pressure, and center image."),
    ],
    "faq": [
        ("What does AuAg mean?", "Au is gold, Ag is silver. On this launch edition, the identity is split left/right: gold and silver accents over hand-finished shell work, with forged-carbon faceplates carrying the visual identity."),
        ("Is this only a cosmetic edition?", "No. The finish attracts attention first, but the reason to stay is the Z05 platform: dynamic bass, BA clarity, Bass-Flow pressure control, and TwinCore voicing."),
        ("What does TwinCore change?", "It changes the overall voicing. Blue leans fuller and easier; red leans cleaner and more awake. The result feels broader than a simple bass bump."),
        ("Why do tips matter?", "With a dynamic woofer, seal is part of the tuning. The right tip helps stabilize bass weight, vocal center, and left-right balance."),
        ("Who is the target listener?", "Listeners who change albums and sources often, want one compact universal monitor for daily use, and value both physical bass and hybrid detail."),
    ],
    "cta": "Choose the mood, keep the groove.",
}


ZH = {
    "cover_kicker": "AAW Z 系列通用型入耳式耳机 - 代理商产品资料",
    "launch": "AuAg 首发版",
    "intro": "一款为不断切换音乐、音源与听感心境的用户打造的紧凑型通用圈铁混合耳机。Z05 AuAg 结合 10mm 动圈冲击力、动铁清晰度、TwinCore 实体调音、手工精饰 AuAg 壳体，以及适合日常轮换的线材套装。",
    "slogan": "善听者，终能找到道路。",
    "position": [
        "Z05 AuAg 的定位是一款小巧但认真的通用贴合圈铁混合耳机：便携、易驱动、外观辨识度强，并且能适应不断变化的歌单与心情。",
        "它不是只为一首试音曲服务的单一听感耳机，而是为真实使用场景打造：整张专辑、播放列表、手机小尾巴、数字音频播放器、紧凑型解码耳放、桌面系统，以及随时变化的聆听心境。",
        "首发版采用 AuAg 视觉身份：Au 是金，Ag 是银。金银色装饰点缀在手工完成的壳体上，由锻造碳纤维面板承担主要视觉重心。"
    ],
    "selling": [
        ("日常通用圈铁混合耳机", "紧凑通用腔体、绕耳式走线，并通过 LiquidMorph 耳塞套获得可重复的密封与舒适度。"),
        ("TwinCore 实体调音", "壳体上的双档实体开关可切换整体声音风格，无需工具、无需 App，也不是难操作的隐藏凹槽。"),
        ("动圈低频骨架", "10mm PET 振膜动圈低音单元让底鼓和贝斯线条拥有真实的压力、质感和身体感。"),
        ("动铁清晰度", "中频与极高频动铁单元提供人声轮廓、镲片空气感和空间信息，同时不削弱声音厚度。"),
        ("信号链一步到位", "2-pin 可换线，包装内附 3.5mm 单端、4.4mm 平衡和 Type-C 可换插头。"),
    ],
    "twin": [
        "蓝色模式", "厚重感 / 流畅感", "更丰满的低频、更柔和的高频谐波，以及更适合长时间聆听整张专辑的中心结像。",
        "红色模式", "聚焦感 / 闪耀感", "更紧致的低频轮廓、更靠前的清晰度，以及当你想让混音打开时额外的镲片空气感。"
    ],
    "specs": [
        ("产品型号", "AAW Z05 AuAg"),
        ("类型", "通用贴合圈铁混合入耳式监听耳机"),
        ("低频单元", "10mm PET 振膜动圈低音单元"),
        ("动铁系统", "中频与极高频动铁单元，提供人声轮廓与空气感"),
        ("声音控制", "TwinCore 实体调音开关"),
        ("压力管理", "Bass-Flow 泄压气流系统"),
        ("壳体", "3D 打印通用监听耳机壳体"),
        ("面板", "锻造碳纤维面板，带 Au/Ag 箔视觉身份"),
        ("饰面", "手工饰以金箔与银箔；每对耗时手工完成"),
        ("耳机端连接", "2-pin 可换线接口"),
        ("音源端插头", "3.5mm 单端、4.4mm 平衡、Type-C"),
        ("灵敏度", "110dB @ 1kHz"),
        ("阻抗", "12 Ohm @ 1kHz"),
        ("总谐波失真", "<0.5% @ 1kHz"),
    ],
    "box": [
        ("Z05 AuAg 耳机本体", "左右通用贴合圈铁混合耳机，采用 AuAg 首发版饰面。"),
        ("2-pin 可换线", "编织线材，耳机端采用常见 2-pin 连接方式，便于长期使用与替换。"),
        ("3.5mm 单端插头", "适用于常见模拟输出与广泛设备兼容。"),
        ("4.4mm 平衡插头", "适用于数字音频播放器与带平衡输出的紧凑型解码耳放。"),
        ("Type-C 插头", "适用于兼容的手机、平板、笔记本电脑与紧凑数字音频场景。"),
        ("收纳盒", "便于存放与携带。"),
        ("LiquidMorph 耳塞套", "透明伞边与青蓝内管，有助于稳定密封、低频压力与中心结像。"),
    ],
    "faq": [
        ("AuAg 是什么意思？", "Au 是金，Ag 是银。在这款首发版上，标识分为左右：在手工饰面的壳体上点缀金银色装饰，由锻造碳纤维面板承担视觉重心。"),
        ("这只是一个换壳版本吗？", "不是。饰面首先吸引目光，但让用户留下来的理由是 Z05 的平台：动圈低频、动铁清晰度、Bass-Flow 压力控制，以及 TwinCore 调音。"),
        ("TwinCore 开关改变了什么？", "它改变整体声音风格。蓝色倾向于更丰满、更轻松；红色倾向于更干净、更精神。它比简单的低频增强更全面。"),
        ("为什么耳塞套很重要？", "对于动圈低音单元，密封就是调音的一部分。正确耳塞套有助于稳定低频厚度、人声中心结像与左右平衡。"),
        ("目标用户是谁？", "经常切换专辑与音源、希望一副紧凑通用耳机覆盖日常使用，并同时重视动圈低频与混合单元细节的听众。"),
    ],
    "cta": "选择心境，留住律动。",
}


def build(data, out_path, zh=False):
    doc = Document()
    setup(doc, zh)
    cover(doc, data, zh)

    kicker(doc, "Positioning" if not zh else "产品定位")
    heading(doc, "Dealer-facing product narrative" if not zh else "供代理商参考的产品叙事")
    for para in data["position"]:
        body(doc, para)
    image(doc, "z05-auag-forged-carbon-faceplate-hero.jpg", 6.3)

    kicker(doc, "Core selling points" if not zh else "核心销售要点")
    heading(doc, "What dealers should lead with" if not zh else "建议代理商优先传达的内容")
    kv_table(doc, data["selling"])

    doc.add_section(WD_SECTION.NEW_PAGE)
    kicker(doc, "TwinCore" if not zh else "TwinCore 双核调音开关")
    heading(doc, "Two listening moods, one physical switch" if not zh else "同一副耳机，两种聆听心境")
    if zh:
        body(doc, "TwinCore 是 Z05 AuAg 与普通单一听感耳机拉开差异的关键。它通过壳体上的实体开关，让用户在更厚重流畅与更清晰开阔之间快速切换。")
    else:
        body(doc, "TwinCore is the feature that separates Z05 AuAg from a single-signature universal monitor. It lets the listener move between a fuller, easier presentation and a cleaner, more open presentation through a real switch on the shell.")
    image(doc, "z05-auag-twincore-switch-closeup.jpg", 6.1)
    matrix_table(
        doc,
        ["Mode", "Character", "Dealer explanation"] if not zh else ["模式", "声音性格", "代理商说明"],
        [
            [data["twin"][0], data["twin"][1], data["twin"][2]],
            [data["twin"][3], data["twin"][4], data["twin"][5]],
        ],
        widths=[1.3, 1.7, 4.0],
    )

    kicker(doc, "Acoustic platform" if not zh else "声学平台")
    heading(doc, "Dynamic bass with hybrid detail" if not zh else "动圈低频与混合单元细节")
    if zh:
        body(doc, "Z05 AuAg 使用 10mm PET 振膜动圈低音单元作为低频基础，并叠加中频与极高频动铁单元，用于补充人声边缘、镲片空气感与空间信息。Bass-Flow 泄压气流系统帮助管理动态低频的压力，让低频更稳定，也减少壳体位移对听感的影响。")
    else:
        body(doc, "Z05 AuAg uses a 10mm PET diaphragm dynamic woofer as its low-frequency foundation, then layers mid and super-high BA drivers for vocal edge, cymbal air, and room cues. Bass-Flow vent pressure management helps stabilize dynamic-driver pressure and reduces the way fit movement affects the listening balance.")
    image(doc, "z05-auag-updated-spec-card.png", 6.7)

    doc.add_section(WD_SECTION.NEW_PAGE)
    kicker(doc, "Craft, fit, and finish" if not zh else "工艺、佩戴与饰面")
    heading(doc, "Made to be worn, then looked at twice" if not zh else "为佩戴而生，随后才被再三审视")
    if zh:
        body(doc, "壳体必须先让人感觉不到它的存在，然后饰面才有机会炫耀。Z05 AuAg 从紧凑的通用贴合、可重复的密封，以及能在长时间聆听中保持说服力的单元布局开始。锻造碳纤维面板承载金/银身份标识，3D 打印壳体再由手工点缀金银箔。")
    else:
        body(doc, "The shell has to disappear before the finish gets to show off. Z05 AuAg starts with a compact universal fit, repeatable seal, and a driver layout that can stay convincing across long sessions. Forged-carbon faceplates carry the gold/silver identity, while the 3D-printed shell receives hand-laced foil work.")
    image(doc, "z05-auag-male-fitment.jpg", 6.3)
    image(doc, "z05-auag-luxury-shell-geometry.jpg", 6.3)

    kicker(doc, "Complete package" if not zh else "包装清单")
    heading(doc, "Accessory and source-path reference" if not zh else "附件与音源路径参考")
    image(doc, "z05-auag-complete-package-flatlay.jpg", 6.3)
    kv_table(doc, data["box"])

    doc.add_section(WD_SECTION.NEW_PAGE)
    kicker(doc, "Technical reference" if not zh else "技术参数参考")
    heading(doc, "Product details and specifications" if not zh else "产品细节与参数")
    kv_table(doc, data["specs"])

    kicker(doc, "Fit and eartips" if not zh else "佩戴与耳塞套")
    heading(doc, "LiquidMorph is part of the tuning" if not zh else "LiquidMorph 耳塞套是调音的一部分")
    if zh:
        body(doc, "对于动圈低音单元，密封绝非细枝末节。耳塞套贴合到位后，低频厚重感、人声中心结像与左右平衡才会稳定。代理商试听时应准备不同尺码耳塞套，让用户先获得正确隔音，再判断声音。")
    else:
        body(doc, "With a dynamic woofer, seal is not a footnote. When the tip sits right, bass weight, vocal center, and left-right balance become stable. Dealers should fit the listener with the right LiquidMorph size before evaluating the sound.")
    image(doc, "z05-liquidmorph-clear-eartips-closeup.jpg", 5.8)
    image(doc, "z05-auag-liquidmorph-accessory-tray.jpg", 5.8)

    kicker(doc, "Dealer FAQ" if not zh else "代理商 FAQ")
    heading(doc, "Reference answers for product conversations" if not zh else "用于销售沟通的参考回答")
    kv_table(doc, data["faq"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run(p, data["cta"], bold=True, size=20, color=BLACK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, data["slogan"], size=9.5, color=MUTED)

    doc.save(out_path)


if __name__ == "__main__":
    build(EN, DOCS / "Z05 AuAg Dealer Product Literature - EN.docx", zh=False)
    build(ZH, DOCS / "Z05 AuAg Dealer Product Literature - 中文.docx", zh=True)
