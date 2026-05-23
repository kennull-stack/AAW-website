from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "aaw-me-update" / "docs"
ASSETS = ROOT / "aaw-theme" / "assets"


COLORS = {
    "black": RGBColor(12, 12, 12),
    "soft_black": RGBColor(28, 28, 28),
    "gold": RGBColor(179, 127, 44),
    "cream": RGBColor(246, 242, 234),
    "body": RGBColor(38, 34, 29),
    "muted": RGBColor(100, 92, 82),
    "white": RGBColor(255, 255, 255),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9D1C4", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=140, start=160, bottom=140, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, bold=False, italic=False, size=None, color=None, all_caps=False):
    run = paragraph.add_run(text.upper() if all_caps else text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=8.5, color=COLORS["gold"], all_caps=True)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = add_run(p, text, bold=True, size=24 if level == 1 else 15, color=COLORS["black"])
    return p


def add_body(doc, text, after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    add_run(p, text, size=9.5, color=COLORS["body"])
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    add_run(p, text, size=9.3, color=COLORS["body"])
    return p


def add_image(doc, name, width=6.6):
    path = ASSETS / name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(7)
        p.add_run().add_picture(str(path), width=Inches(width))


def style_doc(doc, lang):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if lang == "zh" else "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = COLORS["body"]

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei" if lang == "zh" else "Arial")


def add_cover(doc, content, lang):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 10440)
    set_col_widths(table, [Inches(3.3), Inches(3.95)])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "080808")
        set_cell_border(cell, "080808", "0")
        set_cell_margins(cell, 260, 250, 260, 250)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    add_run(p, content["eyebrow"], bold=True, size=8, color=COLORS["gold"], all_caps=lang == "en")
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Z05 AuAg", bold=True, size=30, color=COLORS["white"])
    p = left.add_paragraph()
    add_run(p, content["launch"], bold=True, size=18, color=COLORS["gold"])
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run(p, content["dek"], size=9.5, color=RGBColor(220, 214, 204))
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_run(p, content["slogan"], bold=True, size=9.5, color=COLORS["cream"])

    img_path = ASSETS / "z05-auag-luxury-portrait-hero.png"
    if img_path.exists():
        pr = right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr.add_run().add_picture(str(img_path), width=Inches(3.75))


def add_feature_table(doc, rows, dark=False):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 10200)
    set_col_widths(table, [Inches(1.7), Inches(5.25)])
    for idx, (label, value) in enumerate(rows):
        c0, c1 = table.rows[idx].cells
        for cell in (c0, c1):
            set_cell_border(cell, "D8D0C3" if not dark else "404040", "5")
            set_cell_margins(cell, 130, 160, 130, 160)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "FFFFFF" if not dark else "121212")
        c0.paragraphs[0].add_run(label).bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = COLORS["gold"]
        c0.paragraphs[0].runs[0].font.size = Pt(8.5)
        c1.paragraphs[0].add_run(value)
        c1.paragraphs[0].runs[0].font.color.rgb = COLORS["body"] if not dark else COLORS["white"]
        c1.paragraphs[0].runs[0].font.size = Pt(9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_cards(doc, cards):
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 10200)
    width = Inches(7.05 / len(cards))
    set_col_widths(table, [width] * len(cards))
    for idx, card in enumerate(cards):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "111111")
        set_cell_border(cell, "2C2C2C", "5")
        set_cell_margins(cell, 160, 150, 160, 150)
        p = cell.paragraphs[0]
        add_run(p, card[0], bold=True, size=10.5, color=COLORS["white"])
        p = cell.add_paragraph()
        add_run(p, card[1], size=8.4, color=RGBColor(214, 206, 194))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_faq(doc, faqs):
    for q, a in faqs:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_width(table, 10200)
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, "FFFFFF")
        set_cell_border(cell, "D8D0C3", "5")
        set_cell_margins(cell, 140, 170, 140, 170)
        p = cell.paragraphs[0]
        add_run(p, q, bold=True, size=9.7, color=COLORS["black"])
        p = cell.add_paragraph()
        add_run(p, a, size=9.1, color=COLORS["body"])
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


EN = {
    "eyebrow": "AAW Z Series Universal - AuAg Launch Edition",
    "launch": "Launch Edition",
    "dek": "Z05 AuAg is the small universal you reach for when the playlist keeps changing. It keeps the 10mm dynamic-driver punch, adds armature clarity, and lets TwinCore shift the mood without swapping IEMs.",
    "slogan": "Those who listen will find the way.",
    "sections": [
        ("TwinCore", "Same monitor, two listening moods", [
            "Some albums want warmth and weight. Others ask for cleaner edges and a little more air. TwinCore lets Z05 AuAg move between those moods with a real switch on the shell, large enough to nudge by fingernail while the monitors stay in your ears.",
            "No tool, no tiny recessed slot, no guessing through a phone app. The action has just enough resistance, only two positions, and is easy to find by touch once you know where it sits."
        ]),
        ("Overview", "A compact hybrid for people who rotate music", [
            "Some IEMs are tuned for one perfect demo track. Z05 AuAg is built for the way most of us listen: albums, playlists, dongles, DAPs, desk chains, and a mood that changes by the hour."
        ]),
        ("Fitment", "A universal fit that still feels settled", [
            "The best tuning falls apart if the shell keeps shifting. Z05 AuAg keeps the body compact, routes the cable over-ear, and lets the tip do the sealing work."
        ]),
        ("Signal Chain Ready", "One cable, the sources most of us actually use", [
            "No one wants to buy a new cable on day one. Z05 AuAg ships with the paths that cover the usual rotation: phone, dongle, DAP, and desktop stack."
        ]),
        ("Seal & Comfort", "The tip is part of the tuning", [
            "With a dynamic woofer, seal is not a footnote. When the tip sits right, bass weight, vocal center, and left-right balance stop changing every time you move."
        ]),
    ],
    "cards": [
        ("Dynamic Bass", "Kick drums and bass lines have real skin and pressure, not just a lifted low-end shelf."),
        ("Hybrid Detail", "Armatures bring vocal edges, cymbal air, and room cues without thinning out the body."),
        ("AuAg Finish", "Forged-carbon faceplates carry the gold/silver identity; the shells are finished by hand, pair by pair."),
    ],
    "modes": [
        ("Blue Mode", "Weight / Flow", "A fuller low end, calmer upper harmonics, and a center image that feels easier to sit with for long albums."),
        ("Red Mode", "Focus / Spark", "Tighter bass outlines, more forward clarity, and extra cymbal air when you want the mix to open up."),
    ],
    "specs": [
        ("Format", "Universal-fit hybrid in-ear monitor"),
        ("Low End", "10mm PET diaphragm dynamic woofer"),
        ("Clarity", "Mid and super-high BA drivers for vocal edge and air"),
        ("Voice Control", "TwinCore physical voicing switch"),
        ("Pressure Flow", "Bass-Flow vent pressure system"),
        ("Finish", "Forged-carbon faceplates; 3D-printed shells with hand-laced gold and silver foil finish"),
        ("In The Box", "2-pin detachable cable, 3.5mm, 4.4mm balanced, Type-C terminations, carrying case, LiquidMorph eartips"),
    ],
    "faq": [
        ("What does AuAg mean?", "Au is gold, Ag is silver. On this launch edition, the identity is split left/right: gold and silver accents over the hand-finished shell work, with forged-carbon faceplates doing the visual heavy lifting."),
        ("Is this only a cosmetic edition?", "No. The finish gets attention first, but the reason to stay is the Z05 platform: dynamic bass, BA clarity, Bass-Flow pressure control, and TwinCore voicing."),
        ("What does the TwinCore switch change?", "It changes the overall voicing. Blue leans fuller and easier; red leans cleaner and more awake. It feels broader than a simple bass bump."),
        ("What cable is included?", "Z05 AuAg includes a 2-pin detachable cable, 3.5mm stereo, 4.4mm balanced, and Type-C terminations, plus LiquidMorph tips so you can get the seal right from the first session."),
    ],
    "cta": "Choose the mood, keep the groove.",
    "cta_body": "A launch-edition Z05 with real low-end body, armature detail, and a physical switch for the way your listening mood actually changes.",
}


ZH = {
    "eyebrow": "AAW Z 系列通用型入耳式耳机 - AuAg 首发版",
    "launch": "首发版",
    "dek": "Z05 AuAg 是当你的歌单不断切换时，你会伸手去拿的那副耳机。它保留了 10mm 动圈的冲击力，增加了动铁单元的清晰度，并让你无需更换耳机，就能通过 TwinCore 切换听感情绪。",
    "slogan": "善听者，终能找到道路。",
    "sections": [
        ("TwinCore 双核调音开关", "同一副耳机，两种聆听心境", [
            "有些专辑需要温暖和厚重感。有些则需要更清晰的线条和稍多的空气感。TwinCore 通过壳体上的一个实体开关，让 Z05 AuAg 在这些心境之间切换，开关足够大，无需取出耳机，用指甲即可拨动。",
            "无需工具，没有微小的隐藏卡槽，也不用在手机 App 里瞎猜。开关拨动有适度的阻尼感，只有两个档位，一旦你熟悉了它的位置，就能轻松盲操。"
        ]),
        ("概览", "为不断切换音乐的人打造的紧凑型圈铁混合耳机", [
            "有些入耳式耳机是为那一首完美的试音曲而调音的。而 Z05 AuAg 是为我们大多数人的聆听方式打造的：整张专辑、播放列表、小尾巴、数字音频播放器、桌面设备组合，以及随时变化的心情。"
        ]),
        ("佩戴体验", "一种依然安稳贴合的通用佩戴感", [
            "如果壳体不停移位，再好的调音也会大打折扣。Z05 AuAg 保持紧凑的腔体，采用绕耳式走线，让耳塞套负责密封工作。"
        ]),
        ("信号链一步到位", "一条线材，适配我们大多数人真正使用的音源", [
            "没人想在第一天就去买条新线材。Z05 AuAg 随附的插头覆盖了日常轮换场景：手机、小尾巴、数字音频播放器和桌面组合。"
        ]),
        ("密封与舒适", "耳塞套是调音的一部分", [
            "对于动圈低音单元，密封绝非细枝末节。当耳塞套贴合到位，低频的厚重感、人声的结像和左右平衡就不会再随着你的每次移动而变化。"
        ]),
    ],
    "cards": [
        ("动态低频", "底鼓和贝斯线条拥有真实的质感和冲击力，而不仅仅是抬高的低频响应。"),
        ("混合单元细节", "动铁单元带来人声的轮廓、镲片的空气感和空间信息，同时不削弱声音的厚度。"),
        ("AuAg 饰面", "锻造碳纤维面板承载着金/银身份标识；壳体一对一手工完成饰面。"),
    ],
    "modes": [
        ("蓝色模式", "厚重感 / 流畅感", "更丰满的低频，更柔和的高频谐波，以及一个更适合长时间聆听整张专辑的中心结像。"),
        ("红色模式", "聚焦感 / 闪耀感", "更紧致的低频轮廓，更靠前的清晰度，以及当你想让混音打开时额外的镲片空气感。"),
    ],
    "specs": [
        ("类型", "通用贴合圈铁混合入耳式监听耳机"),
        ("低频", "10mm PET 振膜动圈低音单元"),
        ("清晰度", "中频与极高频动铁单元，提供人声轮廓与空气感"),
        ("声音控制", "TwinCore 实体调音开关"),
        ("气流压力", "Bass-Flow 泄压气流系统"),
        ("饰面", "锻造碳纤维面板；3D 打印壳体，手工饰以金箔与银箔"),
        ("包装清单", "2-pin 可换线、3.5mm 单端插头、4.4mm 平衡插头、Type-C 插头、收纳盒、LiquidMorph 耳塞套"),
    ],
    "faq": [
        ("AuAg 是什么意思？", "Au 是金，Ag 是银。在这款首发版上，标识分为左右：在手工饰面的壳体上点缀金银色装饰，由锻造碳纤维面板承担视觉重心。"),
        ("这只是一个换壳的版本吗？", "不是。饰面首先吸引目光，但让你留下来的理由是 Z05 的平台：动圈低频、动铁清晰度、Bass-Flow 压力控制，以及 TwinCore 调音。"),
        ("TwinCore 开关改变了什么？", "它改变了整体的声音风格。蓝色倾向于更丰满、更轻松；红色倾向于更干净、更精神。感觉它比简单的低频增强要更全面。"),
        ("包含什么样的线材？", "Z05 AuAg 附带一条 2-pin 可换线材、3.5mm 立体声、4.4mm 平衡和 Type-C 可换插头，外加 LiquidMorph 耳塞套，让你从第一次聆听起就能获得正确隔音。"),
    ],
    "cta": "选择心境，留住律动。",
    "cta_body": "一款首发版 Z05，拥有真实的低频厚度、动铁细节，以及一个为你聆听心境变化而设的实体开关。",
}


def build(content, lang, out_path):
    doc = Document()
    style_doc(doc, lang)
    add_cover(doc, content, lang)

    add_kicker(doc, "At a glance" if lang == "en" else "一目了然")
    add_heading(doc, "The build behind the listen" if lang == "en" else "聆听背后的构造")
    add_cards(doc, content["cards"])

    for idx, (kicker, heading, paragraphs) in enumerate(content["sections"]):
        if idx == 2:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_kicker(doc, kicker)
        add_heading(doc, heading)
        for para in paragraphs:
            add_body(doc, para)
        if idx == 0:
            add_image(doc, "z05-auag-twincore-switch-closeup.jpg", 5.8)
            add_feature_table(doc, [(m[0] + " - " + m[1], m[2]) for m in content["modes"]])
        elif idx == 2:
            add_image(doc, "z05-auag-male-fitment.jpg", 6.3)
        elif idx == 3:
            add_image(doc, "z05-auag-complete-package-flatlay.jpg", 6.3)
        elif idx == 4:
            add_image(doc, "z05-liquidmorph-clear-eartips-closeup.jpg", 5.8)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_kicker(doc, "Craft & Sound" if lang == "en" else "工艺与声音")
    add_heading(doc, "Made to be worn, then looked at twice" if lang == "en" else "为佩戴而生，随后才被再三审视")
    add_body(doc, "The shell has to disappear before the finish gets to show off. Z05 AuAg starts with a compact universal fit, repeatable seal, and a driver layout that can stay convincing across long sessions." if lang == "en" else "壳体必须先让人感觉不到它的存在，然后饰面才有机会炫耀。Z05 AuAg 从紧凑的通用贴合、可重复的密封，以及能在长时间聆听中保持说服力的单元布局开始。")
    add_image(doc, "z05-auag-luxury-shell-geometry.jpg", 6.3)
    add_feature_table(doc, content["specs"], dark=False)

    add_kicker(doc, "FAQ" if lang == "en" else "常见问题")
    add_heading(doc, "A few notes before you hit play" if lang == "en" else "在你按下播放键前的几点说明")
    add_faq(doc, content["faq"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, content["cta"], bold=True, size=20, color=COLORS["black"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, content["slogan"], size=9.5, color=COLORS["muted"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, content["cta_body"], size=9.5, color=COLORS["body"])

    doc.save(out_path)


if __name__ == "__main__":
    build(EN, "en", DOCS / "Z05 AuAg Product Literature - EN Updated.docx")
    build(ZH, "zh", DOCS / "Z05 AuAg Product Literature - 中文 Updated.docx")
